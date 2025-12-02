"""
文件解析器 - 支持多种文档格式
"""
import hashlib
from pathlib import Path
from typing import Optional
import chardet


class FileParser:
    """文件解析器"""
    
    SUPPORTED_EXTENSIONS = {
        ".txt": "text",
        ".md": "markdown",
        ".markdown": "markdown",
        ".pdf": "pdf",
        ".docx": "docx",
        ".doc": "doc",
    }
    
    @classmethod
    def get_file_type(cls, filename: str) -> Optional[str]:
        """获取文件类型"""
        ext = Path(filename).suffix.lower()
        return cls.SUPPORTED_EXTENSIONS.get(ext)
    
    @classmethod
    def is_supported(cls, filename: str) -> bool:
        """检查文件是否支持"""
        return cls.get_file_type(filename) is not None
    
    @classmethod
    def compute_hash(cls, content: bytes) -> str:
        """计算内容哈希"""
        return hashlib.sha256(content).hexdigest()
    
    @classmethod
    async def parse(cls, file_path: Path) -> str:
        """
        解析文件内容
        """
        file_type = cls.get_file_type(file_path.name)
        
        if file_type in ("text", "markdown"):
            return await cls._parse_text(file_path)
        elif file_type == "pdf":
            return await cls._parse_pdf(file_path)
        elif file_type == "docx":
            return await cls._parse_docx(file_path)
        elif file_type == "doc":
            return await cls._parse_doc(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    @classmethod
    async def _parse_text(cls, file_path: Path) -> str:
        """解析纯文本文件"""
        raw_content = file_path.read_bytes()
        
        # 检测编码
        detected = chardet.detect(raw_content)
        encoding = detected.get("encoding", "utf-8") or "utf-8"
        
        try:
            return raw_content.decode(encoding)
        except UnicodeDecodeError:
            # 尝试常见编码
            for enc in ["utf-8", "gbk", "gb2312", "latin-1"]:
                try:
                    return raw_content.decode(enc)
                except UnicodeDecodeError:
                    continue
            raise ValueError("Unable to decode file content")
    
    @classmethod
    async def _parse_pdf(cls, file_path: Path) -> str:
        """解析PDF文件"""
        try:
            from pypdf import PdfReader
            
            reader = PdfReader(str(file_path))
            text_parts = []
            
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            
            content = "\n\n".join(text_parts)
            
            # 如果提取的文本太少，可能是扫描版PDF，尝试OCR
            if len(content.strip()) < 100:
                ocr_content = await cls._ocr_pdf(file_path)
                if ocr_content:
                    content = ocr_content
            
            return content
        except ImportError:
            raise ImportError("pypdf is required for PDF parsing")
    
    @classmethod
    async def _ocr_pdf(cls, file_path: Path) -> Optional[str]:
        """使用OCR处理扫描版PDF"""
        try:
            from pdf2image import convert_from_path
            import pytesseract
            
            # 将PDF转换为图片
            images = convert_from_path(str(file_path))
            
            text_parts = []
            for image in images:
                text = pytesseract.image_to_string(image, lang="chi_sim+eng")
                if text:
                    text_parts.append(text)
            
            return "\n\n".join(text_parts)
        except ImportError:
            # OCR依赖未安装，返回None
            return None
        except Exception:
            return None
    
    @classmethod
    async def _parse_docx(cls, file_path: Path) -> str:
        """解析DOCX文件"""
        try:
            from docx import Document
            
            doc = Document(str(file_path))
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # 也提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            return "\n\n".join(text_parts)
        except ImportError:
            raise ImportError("python-docx is required for DOCX parsing")
    
    @classmethod
    async def _parse_doc(cls, file_path: Path) -> str:
        """解析旧版DOC文件"""
        # 旧版.doc格式处理较复杂，建议用户转换为.docx
        raise ValueError("Old .doc format is not supported. Please convert to .docx")


